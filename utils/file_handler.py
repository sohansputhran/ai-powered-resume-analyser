import streamlit as st
import PyPDF2
import docx
import io
from typing import Optional, Union
import tempfile
import os
from config.settings import FILE_CONFIG, ERROR_MESSAGES

class FileHandler:
    """Handle file uploads and text extraction"""
    
    def __init__(self):
        """Initialize file handler with configuration"""
        self.config = FILE_CONFIG
        self.max_file_size = self.config["max_pages"]
        self.supported_formats = self.config["supported_formats"]
    
    def validate_file(self, uploaded_file) -> tuple[bool, str]:
        """
        Validate uploaded file for size and format
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if uploaded_file is None:
            return False, "No file uploaded"
        
        # Check file size
        file_size = uploaded_file.size
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension not in self.supported_formats:
            return False, ERROR_MESSAGES["invalid_format"]
        
        max_size = self.supported_formats[file_extension]["max_size"]
        if file_size > max_size:
            max_size_mb = max_size / (1024 * 1024)
            return False, ERROR_MESSAGES["file_too_large"].format(max_size=max_size_mb)
        
        return True, "File is valid"
    
    def extract_text_from_file(self, uploaded_file) -> str:
        """
        Extract text from uploaded file based on file type
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text content
        """
        # Validate file first
        is_valid, error_message = self.validate_file(uploaded_file)
        if not is_valid:
            st.error(error_message)
            return ""
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        try:
            if file_extension == 'pdf':
                return self._extract_from_pdf(uploaded_file)
            elif file_extension == 'docx':
                return self._extract_from_docx(uploaded_file)
            elif file_extension == 'txt':
                return self._extract_from_txt(uploaded_file)
            else:
                st.error(f"Unsupported file format: {file_extension}")
                return ""
                
        except Exception as e:
            st.error(f"Error extracting text from {uploaded_file.name}: {str(e)}")
            return ""
    
    def _extract_from_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file"""
        try:
            # Read PDF content
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            text_content = []
            
            # Limit pages to prevent excessive processing
            max_pages = min(len(pdf_reader.pages), self.config["max_pages"])
            
            for page_num in range(max_pages):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
            
            extracted_text = "\n".join(text_content).strip()
            
            if not extracted_text:
                st.warning("No text found in PDF. The file might be image-based or corrupted.")
                return ""
            
            return extracted_text
            
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    def _extract_from_docx(self, uploaded_file) -> str:
        """Extract text from DOCX file"""
        try:
            # Create temporary file for docx processing
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_file_path = tmp_file.name
            
            try:
                # Read DOCX content
                doc = docx.Document(tmp_file_path)
                text_content = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                extracted_text = "\n".join(text_content).strip()
                
                if not extracted_text:
                    st.warning("No text found in DOCX file.")
                    return ""
                
                return extracted_text
                
            finally:
                # Clean up temporary file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                    
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    def _extract_from_txt(self, uploaded_file) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)  # Reset file pointer
                    content = uploaded_file.read()
                    
                    if isinstance(content, bytes):
                        text_content = content.decode(encoding)
                    else:
                        text_content = content
                    
                    if text_content.strip():
                        return text_content.strip()
                        
                except UnicodeDecodeError:
                    continue
            
            st.error("Could not decode text file with any supported encoding")
            return ""
            
        except Exception as e:
            st.error(f"Error reading TXT file: {str(e)}")
            return ""
    
    def save_file_temporarily(self, uploaded_file, suffix: str = None) -> str:
        """
        Save uploaded file to temporary location
        
        Args:
            uploaded_file: Streamlit uploaded file object
            suffix: File suffix/extension
            
        Returns:
            Path to temporary file
        """
        try:
            if suffix is None:
                suffix = f".{uploaded_file.name.split('.')[-1]}"
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(uploaded_file.read())
                return tmp_file.name
                
        except Exception as e:
            st.error(f"Error saving temporary file: {str(e)}")
            return ""
    
    def cleanup_temp_file(self, file_path: str):
        """
        Clean up temporary file
        
        Args:
            file_path: Path to temporary file to delete
        """
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
        except Exception as e:
            st.warning(f"Could not clean up temporary file: {str(e)}")
    
    def get_file_info(self, uploaded_file) -> dict:
        """
        Get information about uploaded file
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dictionary with file information
        """
        if uploaded_file is None:
            return {}
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        file_size_mb = uploaded_file.size / (1024 * 1024)
        
        return {
            "name": uploaded_file.name,
            "size_bytes": uploaded_file.size,
            "size_mb": round(file_size_mb, 2),
            "extension": file_extension,
            "mime_type": uploaded_file.type,
            "is_valid": self.validate_file(uploaded_file)[0]
        }